import os
import re
import pandas as pd
import requests
import json
import logging
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from langchain_community.document_loaders import PyPDFLoader
from paddleocr import PaddleOCR
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document as LangchainDocument

# 屏蔽 PaddleOCR 日志
logging.getLogger('ppocr').setLevel(logging.WARNING)


# ====================== 【全局配置】 ======================
# RAG 分块配置
CHUNK_SIZE = 512
CHUNK_OVERLAP = 64
MIN_CHUNK_SIZE = 64

# Ollama 本地模型（你拉的什么就写什么）
OLLAMA_MODEL = "qwen3:4b"

# 向量库路径
VECTOR_DB_PATH = "./chroma_local_db"

# OCR 初始化（图片识别）
ocr = PaddleOCR(use_textline_orientation=True, lang="ch")

# ====================== 【工具：中文语义分块】 ======================
def split_text_by_sentence(text: str):
    sentences = re.split(r'([。！？；;\n])', text)
    sentences = [s1+s2 for s1,s2 in zip(sentences[0::2], sentences[1::2])] if len(sentences)>1 else [text]
    sentences = [s.strip() for s in sentences if s.strip()]

    chunks = []
    current = []
    count = 0

    for s in sentences:
        s_len = len(s)
        if count + s_len > CHUNK_SIZE:
            if current:
                chunk = ''.join(current)
                chunks.append(chunk)
                overlap = chunk[-CHUNK_OVERLAP:] if len(chunk)>=CHUNK_OVERLAP else ''
                current = [overlap, s]
                count = len(overlap) + s_len
        else:
            current.append(s)
            count += s_len

    if current:
        chunk = ''.join(current).strip()
        if len(chunk) >= MIN_CHUNK_SIZE:
            chunks.append(chunk)
    return chunks

# ====================== 【多格式文件解析】 ======================
def parse_txt(filepath):
    with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def parse_docx(filepath):
    doc = Document(filepath)
    return '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

def parse_pdf(filepath):
    loader = PyPDFLoader(filepath)
    pages = loader.load_and_split()
    return '\n'.join([p.page_content for p in pages])

def parse_excel(filepath):
    df = pd.read_excel(filepath)
    lines = []
    for _, row in df.iterrows():
        row_str = ' | '.join([f'{k}:{v}' for k, v in row.to_dict().items() if pd.notna(v)])
        lines.append(row_str)
    return '\n'.join(lines)

def parse_image(filepath):
    result = ocr.ocr(filepath, cls=True)
    lines = []
    for line in result:
        for word in line:
            lines.append(word[1][0])
    return '\n'.join(lines)

def parse_web(url):
    headers = {"User-Agent": "Mozilla/5.0"}
    r = requests.get(url, headers=headers, timeout=10)
    r.encoding = 'utf-8'
    soup = BeautifulSoup(r.text, 'html.parser')
    for tag in soup(['script', 'style', 'nav', 'footer', 'header']):
        tag.decompose()
    return soup.get_text(sep='\n', strip=True)

# 统一解析入口
def parse_file(filepath):
    ext = Path(filepath).suffix.lower()
    try:
        if ext == '.txt':
            text = parse_txt(filepath)
        elif ext == '.docx':
            text = parse_docx(filepath)
        elif ext == '.pdf':
            text = parse_pdf(filepath)
        elif ext in ['.xlsx', '.xls', '.csv']:
            text = parse_excel(filepath)
        elif ext in ['.png', '.jpg', '.jpeg', '.bmp']:
            text = parse_image(filepath)
        else:
            print(f"不支持的格式: {ext}")
            return []
        chunks = split_text_by_sentence(text)
        print(f"✅ {filepath} 完成解析 | 分块数：{len(chunks)}")
        return chunks
    except Exception as e:
        print(f"❌ 解析失败 {filepath}: {e}")
        return []

# ====================== 【构建本地向量库】 ======================
def build_knowledge_base(file_list):
    embedding = HuggingFaceEmbeddings(model_name="BAAI/bge-small-zh-v1.5")

    if os.path.exists(VECTOR_DB_PATH):
        print("🔍 加载已有向量库...")
        db = Chroma(persist_directory=VECTOR_DB_PATH, embedding_function=embedding)
        return db

    print("🚀 新建向量库...")
    documents = []
    for fp in file_list:
        chunks = parse_file(fp)
        for c in chunks:
            documents.append(
                LangchainDocument(page_content=c, metadata={"source": fp})
            )

    db = Chroma.from_documents(
        documents=documents,
        embedding=embedding,
        persist_directory=VECTOR_DB_PATH
    )
    db.persist()
    print("✅ 向量库构建完成！")
    return db

# ====================== 【Ollama 本地问答】 ======================
def chat_with_rag(db, query):
    # 检索相关文档
    docs = db.similarity_search(query, k=3)
    context = '\n\n'.join([d.page_content for d in docs])
    
    # 直接调用 Ollama API
    prompt = f"""基于以下上下文回答问题：

{context}

问题：{query}
"""
    
    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": OLLAMA_MODEL,
            "prompt": prompt,
            "stream": False
        }
    )
    
    answer = response.json().get("response", "")
    return answer, docs


# ====================== 【主程序：运行】 ======================
if __name__ == "__main__":
    # ========== 1. 把你的文件放在这里 ==========
    FILES = [
        "docs/hongkong-gaocai-colleges.pdf",
        "docs/wechat-rpa-arch.png",
        # "docs/data.xlsx",
        # "docs/info.txt",
        # "docs/screenshot.png",
    ]

    # ========== 2. 构建知识库 ==========
    db = build_knowledge_base(FILES)

    # ========== 3. 本地问答 ==========
    print("\n💬 本地 RAG 问答已启动（输入 exit 退出）")
    while True:
        question = input("\n请输入问题：")
        if question.lower() in ["exit", "quit", "q"]:
            break

        answer, sources = chat_with_rag(db, question)
        print("\n🤖 回答：")
        print(answer)

        print("\n📄 参考来源：")
        for s in sources:
            print(f"- {s.metadata['source']}")